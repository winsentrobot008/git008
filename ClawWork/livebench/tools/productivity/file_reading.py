"""
File creation tool supporting multiple formats
"""

from langchain_core.tools import tool
from typing import Dict, Any, List, Union
from pathlib import Path
from typing import Any
from io import BytesIO
import base64
import io
import os
import re
import time
import tempfile
import subprocess
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader, PdfWriter
from typing import Optional

load_dotenv()



# Import global state from parent module
def _get_global_state():
    """Get global state from parent module"""
    from livebench.tools.direct_tools import _global_state
    return _global_state


@tool
def read_file(filetype: str, file_path: Path) -> Dict[str, Any]:
    """
    Read a file and return the content in a format suitable for LLM consumption.

    Supported file types:
    - pdf: PDF documents (converted to images, 4 pages per combined image)
    - docx: Microsoft Word documents
    - xlsx: Microsoft Excel spreadsheets
    - pptx: PowerPoint presentations (converted to slide images)
    - png/jpg/jpeg: Image files (returns base64 for multimodal LLM input)
    - txt: Plain text files

    Args:
        filetype: The type of the file (pdf, docx, xlsx, pptx, png, jpg, jpeg, txt)
        file_path: The path to the file

    Returns:
        Dict with file content. For images/PDFs/PPTX, includes 'images' field with image bytes.
        For text-based files, includes 'text' field with extracted text.
    """
    filetype = filetype.lower().strip()
    
    if filetype == "pdf":
        # Check if model supports multimodal input
        global_state = _get_global_state()
        supports_multimodal = global_state.get("supports_multimodal", True)
        
        if supports_multimodal:
            # Use image-based approach for multimodal models
            print(f"📄 Reading PDF via read_pdf_as_images()")
            images = read_pdf_as_images(file_path)
            if images:
                total_pages = len(images) * 4  # Approximate (last image may have fewer pages)
                return {
                    "type": "pdf_images",
                    "images": images,
                    "image_count": len(images),
                    "approximate_pages": total_pages,
                    "message": f"PDF loaded with ~{total_pages} pages as {len(images)} combined images (4 pages per image). Use images in multimodal LLM calls."
                }
            else:
                raise RuntimeError(
                    "PDF conversion failed. Ensure poppler-utils and pdf2image are installed.\n"
                    "Install with: sudo apt-get install poppler-utils && pip install pdf2image Pillow"
                )
        else:
            # Use OCR-based approach for text-only models
            print(f"📄 Reading PDF via read_pdf_ocr() → _call_qwen_ocr()")
            text = read_pdf_ocr(file_path)
            return {
                "type": "text",
                "text": text,
                "message": f"PDF processed via OCR (model does not support multimodal input)."
            }
    
    elif filetype == "docx":
        print(f"📄 Reading DOCX via read_docx()")
        text = read_docx(file_path)
        return {"type": "text", "text": text}
    
    elif filetype == "xlsx":
        print(f"📊 Reading XLSX via read_xlsx()")
        text = read_xlsx(file_path)
        return {"type": "text", "text": text}
    
    elif filetype == "pptx":
        print(f"📊 Reading PPTX via read_pptx_as_images()")
        images = read_pptx_as_images(file_path)
        if images:
            return {
                "type": "pptx_images",
                "images": images,
                "slide_count": len(images),
                "message": f"PowerPoint presentation loaded with {len(images)} slides. Use images in multimodal LLM calls."
            }
        else:
            raise RuntimeError(
                "PPTX conversion failed. Ensure LibreOffice and pdf2image are installed.\n"
                "Install with: sudo apt-get install libreoffice poppler-utils && pip install pdf2image Pillow"
            )
    
    elif filetype in ["png", "jpg", "jpeg"]:
        print(f"🖼️  Reading {filetype.upper()} via read_image()")
        image_data = read_image(file_path, filetype)
        return {
            "type": "image",
            "image_data": image_data,
            "message": f"Image file loaded. Use this data in multimodal LLM calls with image_url format."
        }
    
    elif filetype == "txt":
        print(f"📝 Reading TXT via read_txt()")
        text = read_txt(file_path)
        return {"type": "text", "text": text}
    
    else:
        raise ValueError(
            f"Unsupported file type: {filetype}. "
            f"Supported types: pdf, docx, xlsx, pptx, png, jpg, jpeg, txt"
        )



def read_docx(docx_path: Path) -> str:
    """
    Read a Microsoft Word document and extract text.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    try:
        doc = Document(str(docx_path))
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_data))
            if table_data:
                tables_text.append("\n".join(table_data))
        
        # Combine all text
        all_text = "\n\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n=== TABLES ===\n\n" + "\n\n".join(tables_text)
        
        return all_text
        
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX file: {str(e)}")


def read_xlsx(xlsx_path: Path) -> str:
    """
    Read a Microsoft Excel spreadsheet and extract data.
    
    Args:
        xlsx_path: Path to XLSX file
        
    Returns:
        Formatted text representation of spreadsheet data
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    
    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"XLSX file not found: {xlsx_path}")
    
    try:
        wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
        
        result = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            result.append(f"=== SHEET: {sheet_name} ===\n")
            
            # Get data from sheet
            data = []
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    data.append(row_str)
            
            if data:
                result.append("\n".join(data))
            else:
                result.append("(Empty sheet)")
            
            result.append("\n")
        
        return "\n".join(result)
        
    except Exception as e:
        raise RuntimeError(f"Failed to read XLSX file: {str(e)}")


def read_image(image_path: Path, filetype: str) -> str:
    """
    Read an image file and return base64-encoded data URL.
    
    This format is compatible with OpenAI's vision API format:
    {"type": "image_url", "image_url": {"url": "data:image/png;base64,..."}}
    
    Args:
        image_path: Path to image file
        filetype: Image type (png, jpg, jpeg)
        
    Returns:
        Base64-encoded data URL string
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image file not found: {image_path}")
    
    try:
        with open(image_path, 'rb') as f:
            image_bytes = f.read()
        
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        # Map filetype to MIME type
        mime_map = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg'
        }
        mime_type = mime_map.get(filetype.lower(), 'image/png')
        
        # Return data URL format
        data_url = f"data:{mime_type};base64,{image_base64}"
        return data_url
        
    except Exception as e:
        raise RuntimeError(f"Failed to read image file: {str(e)}")


def read_txt(txt_path: Path) -> str:
    """
    Read a plain text file.
    
    Args:
        txt_path: Path to TXT file
        
    Returns:
        File content as string
    """
    if not os.path.exists(txt_path):
        raise FileNotFoundError(f"Text file not found: {txt_path}")
    
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to read text file: {str(e)}")


def read_pptx_as_images(pptx_path: Path) -> Optional[List[bytes]]:
    """
    Convert PPTX to list of PNG images (one per slide).
    
    Uses LibreOffice to convert PPTX → PDF, then pdf2image to convert PDF → PNG images.
    
    Args:
        pptx_path: Path to PPTX file
        
    Returns:
        List of PNG image bytes (one per slide), or None if conversion fails
    """
    if not os.path.exists(pptx_path):
        raise FileNotFoundError(f"PPTX file not found: {pptx_path}")
    
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError as e:
        print(f"Required packages not installed: {e}")
        print("Install with: pip install Pillow pdf2image")
        return None
    
    temp_dir = None
    try:
        temp_dir = tempfile.mkdtemp()
        
        # Convert PPTX to PDF using LibreOffice
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, str(pptx_path)],
            capture_output=True,
            timeout=30,
            text=True
        )
        
        if result.returncode != 0:
            print(f"LibreOffice conversion failed: {result.stderr}")
            return None
        
        # Find the generated PDF
        pdf_name = os.path.splitext(os.path.basename(str(pptx_path)))[0] + '.pdf'
        pdf_path = os.path.join(temp_dir, pdf_name)
        
        if not os.path.exists(pdf_path):
            print(f"PDF not found at {pdf_path}")
            return None
        
        # Convert PDF to images (one per slide)
        images = convert_from_path(pdf_path, dpi=150)
        
        # Convert PIL images to PNG bytes
        image_bytes_list = []
        for img in images:
            # Resize to reasonable size (A4 aspect ratio, max 1200px width)
            max_width = 1200
            if img.width > max_width:
                ratio = max_width / img.width
                new_height = int(img.height * ratio)
                img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
            
            # Convert to PNG bytes
            buf = BytesIO()
            img.save(buf, format='PNG')
            image_bytes_list.append(buf.getvalue())
        
        return image_bytes_list
    
    except subprocess.TimeoutExpired:
        print("LibreOffice conversion timed out")
        return None
    except FileNotFoundError:
        print("LibreOffice not found - install with: sudo apt-get install libreoffice")
        return None
    except Exception as e:
        print(f"PPTX to image conversion failed: {str(e)}")
        return None
    finally:
        # Cleanup temp files
        if temp_dir:
            import shutil
            try:
                shutil.rmtree(temp_dir)
            except:
                pass


def read_pdf_as_images(pdf_path: Path) -> Optional[List[bytes]]:
    """
    Convert PDF to list of PNG images, combining 4 pages into one image to save resources.
    
    Uses pdf2image to convert PDF pages to images, then combines them in a 2x2 grid.
    This significantly reduces token usage and API costs.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        List of PNG image bytes (4 pages per image), or None if conversion fails
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError as e:
        print(f"Required packages not installed: {e}")
        print("Install with: pip install Pillow pdf2image")
        return None
    
    try:
        # Convert PDF to images (one per page)
        # Use lower DPI for efficiency (100 instead of 150)
        images = convert_from_path(str(pdf_path), dpi=100)
        
        if not images:
            print(f"No pages found in PDF: {pdf_path}")
            return None
        
        # Combine 4 pages into one image (2x2 grid)
        combined_images = []
        pages_per_image = 4
        
        for i in range(0, len(images), pages_per_image):
            batch = images[i:i + pages_per_image]
            
            # Resize each page to max 600px width (A4 aspect ratio)
            max_width = 600
            resized_batch = []
            
            for img in batch:
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                resized_batch.append(img)
            
            # Calculate grid dimensions (2x2 layout)
            cols = 2
            rows = (len(resized_batch) + cols - 1) // cols  # Ceiling division
            
            # Find max dimensions for uniform grid
            max_page_width = max(img.width for img in resized_batch)
            max_page_height = max(img.height for img in resized_batch)
            
            # Create combined image with white background
            combined_width = max_page_width * cols
            combined_height = max_page_height * rows
            combined = Image.new('RGB', (combined_width, combined_height), 'white')
            
            # Paste each page into the grid
            for idx, img in enumerate(resized_batch):
                row = idx // cols
                col = idx % cols
                x = col * max_page_width
                y = row * max_page_height
                combined.paste(img, (x, y))
            
            # Convert to PNG bytes
            buf = BytesIO()
            combined.save(buf, format='PNG', optimize=True)
            combined_images.append(buf.getvalue())
        
        return combined_images
    
    except FileNotFoundError:
        print("poppler-utils not found - install with: sudo apt-get install poppler-utils")
        return None
    except Exception as e:
        print(f"PDF to image conversion failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def read_pdf_ocr(pdf_path: Path) -> str:
    """
    使用 SiliconFlow DeepSeek-OCR 对单个 PDF 做 OCR。
    
    NOTE: This is the old OCR-based method. The default is now read_pdf_as_images()
    which converts to images for multimodal LLM input.

    输入:
        pdf_path: PDF 文件路径 (Path)

    返回:
        ocr_text: markdown 文本
    """
    raw_result = _call_qwen_ocr(pdf_path)

    ocr_text = raw_result.get("content", "")
    if not isinstance(ocr_text, str):
        ocr_text = str(ocr_text)

    return ocr_text


def _extract_pdf_page_as_pdf(pdf_path: Path, page_num: int) -> bytes:
    """
    从 PDF 中提取指定页面并返回单页 PDF 的字节数据。
    page_num 从 0 开始。
    """
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()

    writer.add_page(reader.pages[page_num])

    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _call_qwen_ocr(pdf_path: Path) -> dict[str, Any]:
    """
    使用 Qwen VL OCR 对多页 PDF 逐页识别并返回原始结构：
    {
        "content": "<text>",
        "total_pages": ...,
        "model": "...",
        "usage": {...},
    }
    
    Optimizations:
    - Limits image resolution to 1536px on longest side for faster processing
    - Uses parallel processing for multiple pages
    
    Cost tracking:
    - OCR pricing: 0.0003 CNY/1K tokens input, 0.0005 CNY/1K tokens output
    - Converted to USD at ~7.2 CNY/USD: $0.0417/1M input, $0.0694/1M output
    """
    if not pdf_path.is_file():
        raise FileNotFoundError(f"找不到 PDF 文件：{pdf_path}")

    api_key = os.getenv("OCR_VLLM_API_KEY")
    if not api_key:
        raise ValueError("未找到 OCR_VLLM_API_KEY 环境变量，请设置 API key")

    # Convert PDF to PNG images (one per page)
    try:
        from pdf2image import convert_from_path
        from PIL import Image
    except ImportError:
        raise ImportError("pdf2image not installed. Run: pip install pdf2image Pillow")
    
    # Convert PDF pages to images
    images = convert_from_path(str(pdf_path), dpi=150)
    total_pages = len(images)
    print(f"total_pages: {total_pages}")
    
    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    # Helper function to process a single page
    def process_page(page_num: int, img: Image.Image) -> tuple[int, str, dict]:
        """Process a single PDF page with OCR"""
        # Resize image to max 1536px on longest side for faster OCR
        max_size = 1536
        if max(img.width, img.height) > max_size:
            if img.width > img.height:
                new_width = max_size
                new_height = int(img.height * (max_size / img.width))
            else:
                new_height = max_size
                new_width = int(img.width * (max_size / img.height))
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Convert PIL image to PNG base64
        buf = BytesIO()
        img.save(buf, format='PNG', optimize=True)
        page_png_base64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{page_png_base64}"
                        },
                    },
                    {
                        "type": "text",
                        "text": "请仅输出图像中的文本内容。",
                    },
                ],
            }
        ]

        response = client.chat.completions.create(
            model="qwen-vl-ocr-2025-11-20",
            messages=messages,
        )

        page_content = response.choices[0].message.content or ""
        usage = {
            "prompt_tokens": response.usage.prompt_tokens if response.usage else 0,
            "completion_tokens": response.usage.completion_tokens if response.usage else 0,
            "total_tokens": response.usage.total_tokens if response.usage else 0,
        }
        
        return (page_num, page_content, usage)

    # Process pages in parallel using ThreadPoolExecutor
    all_pages_content: list[str] = [""] * total_pages  # Pre-allocate list
    total_usage = {
        "prompt_tokens": 0,
        "completion_tokens": 0,
        "total_tokens": 0,
    }

    # Use max 4 workers to avoid overwhelming the API
    max_workers = min(4, total_pages)
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all page processing tasks
        futures = {
            executor.submit(process_page, page_num, img): page_num 
            for page_num, img in enumerate(images)
        }
        
        # Collect results as they complete
        for future in as_completed(futures):
            page_num, page_content, usage = future.result()
            all_pages_content[page_num] = page_content
            
            # Accumulate usage stats
            total_usage["prompt_tokens"] += usage["prompt_tokens"]
            total_usage["completion_tokens"] += usage["completion_tokens"]
            total_usage["total_tokens"] += usage["total_tokens"]
            
            print(f"✓ Page {page_num + 1}/{total_pages} processed")

    full_content = "\n\n".join(all_pages_content)

    # Track OCR API cost
    # OCR pricing: 0.0003 CNY/1K tokens in, 0.0005 CNY/1K tokens out
    # Convert to USD (1 USD ≈ 7.2 CNY): $0.0417/1M in, $0.0694/1M out
    try:
        global_state = _get_global_state()
        tracker = global_state.get("economic_tracker")
        if tracker:
            # Track input tokens cost
            input_cost = tracker.track_api_call(
                tokens=total_usage["prompt_tokens"],
                price_per_1m=0.0417,  # ~0.0003 CNY/1K tokens = $0.0417/1M tokens
                api_name="OCR_Input"
            )
            # Track output tokens cost
            output_cost = tracker.track_api_call(
                tokens=total_usage["completion_tokens"],
                price_per_1m=0.0694,  # ~0.0005 CNY/1K tokens = $0.0694/1M tokens
                api_name="OCR_Output"
            )
            total_ocr_cost = input_cost + output_cost
            print(f"💰 OCR cost: ${total_ocr_cost:.6f} (input: ${input_cost:.6f}, output: ${output_cost:.6f})")
    except Exception as e:
        # Don't fail OCR if cost tracking fails
        print(f"⚠️ Failed to track OCR cost: {e}")

    return {
        "content": full_content,
        "total_pages": total_pages,
        "model": "qwen-vl-ocr-2025-11-20",
        "usage": total_usage,
    }

def _pdf_to_png_base64_list(pdf_path: Path, poppler_path: Optional[str] = None) -> list[str]:
    """
    将 PDF 全部页面转换为 PNG base64 列表，用于 Qwen image_url。
    依赖 pdf2image 和 poppler；若未安装则给出明确错误。
    逐页转换，确保多页文件不会被截断。
    """
    try:
        from pdf2image import convert_from_path
    except Exception as e:
        raise RuntimeError("需要安装 pdf2image 才能将 PDF 转为图片；pip install pdf2image，并确保系统有 poppler") from e

    reader = PdfReader(str(pdf_path))
    total_pages = len(reader.pages)
    if total_pages == 0:
        raise ValueError(f"PDF 无页面或读取失败：{pdf_path}")

    result: list[str] = []
    # 逐页转换，避免只拿到首页
    for page_idx in range(1, total_pages + 1):
        images = convert_from_path(
            str(pdf_path),
            first_page=page_idx,
            last_page=page_idx,
            poppler_path=poppler_path,
        )
        if not images:
            continue
        buf = BytesIO()
        images[0].save(buf, format="PNG")
        png_base64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        result.append(f"data:image/png;base64,{png_base64}")

    if not result:
        raise ValueError(f"PDF 转图片失败：{pdf_path}")

    return result

if __name__ == "__main__":
    pdf = Path(
        "/Users/tianyu/Desktop/workspace/-Live-Bench/livebench/data/test_data/AITrader.pdf"
    )

    result = read_pdf(pdf)
    print(result)