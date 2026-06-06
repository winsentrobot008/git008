"""
LLM-based Work Evaluator using Category-Specific Meta-Prompts

This module implements evaluation using GPT and the comprehensive evaluation
criteria from eval/meta_prompts/ for each task category (occupation).
"""

import os
import json
import base64
from typing import Dict, Optional, Tuple, List, Union
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()


class LLMEvaluator:
    """
    LLM-based evaluator that uses category-specific meta-prompts
    to evaluate agent work artifacts with a 0.0-1.0 score.
    """

    def __init__(
        self,
        meta_prompts_dir: str = "./eval/meta_prompts",
        model: str = "gpt-4o",
        max_payment: float = 50.0
    ):
        """
        Initialize LLM Evaluator

        Args:
            meta_prompts_dir: Path to directory containing evaluation meta-prompts
            model: OpenAI model to use for evaluation
            max_payment: Maximum payment for perfect work
        """
        self.meta_prompts_dir = Path(meta_prompts_dir)
        self.model = model
        self.max_payment = max_payment
        
        # Initialize OpenAI client with separate evaluation configuration
        # Priority: EVALUATION_API_KEY > OPENAI_API_KEY
        api_key = os.getenv("EVALUATION_API_KEY") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("Neither EVALUATION_API_KEY nor OPENAI_API_KEY found in environment")
        
        # Priority: EVALUATION_API_BASE > OPENAI_API_BASE
        base_url = os.getenv("EVALUATION_API_BASE") or os.getenv("OPENAI_API_BASE")
        
        # Allow overriding evaluation model
        if os.getenv("EVALUATION_MODEL"):
            self.model = os.getenv("EVALUATION_MODEL")
        
        # Log configuration
        if os.getenv("EVALUATION_API_KEY"):
            print(f"🔧 Evaluation using separate API key (EVALUATION_API_KEY)")
        else:
            print(f"🔧 Evaluation using shared API key (OPENAI_API_KEY)")
        
        if base_url:
            print(f"🔧 Evaluation API base URL: {base_url}")
            self.client = OpenAI(api_key=api_key, base_url=base_url)
        else:
            print(f"🔧 Evaluation using default OpenAI endpoint")
            self.client = OpenAI(api_key=api_key)
        
        print(f"🔧 Evaluation model: {self.model}")
        
        # Cache for loaded meta-prompts
        self._meta_prompt_cache: Dict[str, Dict] = {}

    def evaluate_artifact(
        self,
        task: Dict,
        artifact_paths: list[str],
        description: str = "",
        max_payment: Optional[float] = None
    ) -> Tuple[float, str, float]:
        """
        Evaluate work artifact(s) using LLM and category-specific criteria

        Args:
            task: Task dictionary with occupation, sector, prompt, etc.
            artifact_paths: List of paths to submitted artifacts
            description: Agent's description of the work
            max_payment: Task-specific max payment (uses default if None)

        Returns:
            Tuple of (evaluation_score 0.0-1.0, feedback_text, payment_amount)
        """
        # Use task-specific max_payment if provided, otherwise fall back to default
        if max_payment is None:
            max_payment = self.max_payment

        # Get task category (occupation)
        occupation = task.get('occupation', '')

        if not occupation:
            return (0.0, "Error: Task missing occupation field", 0.0)

        # Load meta-prompt for this category
        meta_prompt = self._load_meta_prompt(occupation)

        if not meta_prompt:
            # Raise error if no meta-prompt available - no fallback
            raise FileNotFoundError(
                f"No meta-prompt found for occupation '{occupation}'. "
                f"LLM evaluation requires category-specific rubrics. "
                f"Check that eval/meta_prompts/ contains the appropriate file."
            )

        # Check if artifacts exist
        existing_artifacts = []
        missing_artifacts = []

        for path in artifact_paths:
            if os.path.exists(path):
                existing_artifacts.append(path)
            else:
                missing_artifacts.append(path)

        if not existing_artifacts:
            return (
                0.0,
                f"No artifacts found at specified paths: {artifact_paths}",
                0.0
            )

        # Read artifact contents (with size limits for API)
        artifact_data = self._read_artifacts_with_images(existing_artifacts)

        # Build evaluation request with multimodal support
        user_message_content = self._build_multimodal_evaluation_content(
            meta_prompt=meta_prompt,
            task=task,
            artifact_data=artifact_data,
            missing_artifacts=missing_artifacts,
            description=description
        )

        # Call LLM for evaluation
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert work evaluator. Follow the provided rubric precisely and output a structured evaluation."
                    },
                    {
                        "role": "user",
                        "content": user_message_content
                    }
                ],
                # temperature=0.3,  # Lower temperature for consistent evaluation
                # max_tokens=2000
            )

            evaluation_text = response.choices[0].message.content

            # Parse evaluation score from response
            score = self._extract_score(evaluation_text)

            # Convert 0-10 score to 0.0-1.0 scale
            normalized_score = score / 10.0

            # Calculate payment based on score using task-specific max_payment
            payment = normalized_score * max_payment

            return (normalized_score, evaluation_text, payment)

        except Exception as e:
            error_msg = f"LLM evaluation failed: {str(e)}"
            print(f"❌ {error_msg}")

            # Log detailed error information
            from livebench.utils.logger import log_error
            log_error(
                "LLM evaluation API call failed",
                context={
                    "model": self.model,
                    "occupation": occupation,
                    "task_id": task.get('task_id'),
                    "api_base": os.getenv("OPENAI_API_BASE", "default"),
                    "error_type": type(e).__name__,
                    "has_api_key": bool(os.getenv("OPENAI_API_KEY"))
                },
                exception=e
            )

            # Re-raise the error - no fallback
            raise RuntimeError(f"LLM evaluation failed and no fallback is configured: {error_msg}") from e

    def _load_meta_prompt(self, occupation: str) -> Optional[Dict]:
        """
        Load meta-prompt for a specific occupation category

        Args:
            occupation: Occupation name (e.g., "Software_Developers")

        Returns:
            Meta-prompt dictionary or None if not found
        """
        # Normalize occupation name to match file naming
        normalized = occupation.replace(' ', '_').replace(',', '')
        
        # Check cache first
        if normalized in self._meta_prompt_cache:
            return self._meta_prompt_cache[normalized]
        
        # Try to find matching meta-prompt file
        meta_prompt_path = self.meta_prompts_dir / f"{normalized}.json"
        
        if not meta_prompt_path.exists():
            print(f"⚠️ No meta-prompt found for occupation: {occupation}")
            print(f"   Looking for: {meta_prompt_path}")
            return None
        
        # Load and cache
        try:
            with open(meta_prompt_path, 'r', encoding='utf-8') as f:
                meta_prompt = json.load(f)
            
            self._meta_prompt_cache[normalized] = meta_prompt
            return meta_prompt
            
        except Exception as e:
            print(f"⚠️ Error loading meta-prompt for {occupation}: {e}")
            return None

    def _read_artifacts(self, artifact_paths: list[str], max_size_kb: int = 2000) -> Dict[str, str]:
        """
        Read artifact file contents with size limits and format extraction

        Args:
            artifact_paths: List of artifact file paths
            max_size_kb: Maximum file size to read in KB (default 2MB)

        Returns:
            Dictionary mapping file paths to contents
        """
        contents = {}
        
        for path in artifact_paths:
            try:
                file_size = os.path.getsize(path)
                file_ext = os.path.splitext(path)[1].lower()
                
                # Check file size
                if file_size > max_size_kb * 1024:
                    contents[path] = f"[File too large: {file_size} bytes (>{max_size_kb}KB). Only showing metadata]"
                elif file_size == 0:
                    contents[path] = "[Empty file]"
                else:
                    # Handle different file types
                    if file_ext == '.docx':
                        contents[path] = self._read_docx_content(path)
                    elif file_ext == '.xlsx':
                        contents[path] = self._read_xlsx_content(path)
                    elif file_ext in ['.png', '.jpg', '.jpeg', '.gif']:
                        contents[path] = f"[Image file: {file_ext}, {file_size} bytes. Image analysis not yet implemented - evaluator should assess based on task requirements and file existence]"
                    elif file_ext == '.pdf':
                        contents[path] = f"[PDF file: {file_size} bytes. Content extraction not yet implemented - evaluator should assess based on task requirements and file existence]"
                    else:
                        # Try to read as text
                        try:
                            with open(path, 'r', encoding='utf-8') as f:
                                contents[path] = f.read()
                        except UnicodeDecodeError:
                            contents[path] = f"[Binary file: {file_ext}, {file_size} bytes]"
                        
            except Exception as e:
                contents[path] = f"[Error reading file: {str(e)}]"
        
        return contents
    
    def _read_docx_content(self, path: str) -> str:
        """
        Extract text content from DOCX file
        
        Args:
            path: Path to DOCX file
            
        Returns:
            Extracted text content with paragraphs and tables
            
        Raises:
            ImportError: If python-docx is not installed
            RuntimeError: If DOCX extraction fails
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(path)
            
            content = []
            content.append(f"[DOCX Document - {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables]\n")
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                content.append(f"\n--- Table {i+1} ({len(table.rows)} rows × {len(table.columns)} cols) ---")
                for row in table.rows[:10]:  # Limit to first 10 rows
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content.append(row_text)
                if len(table.rows) > 10:
                    content.append(f"... ({len(table.rows) - 10} more rows)")
            
            return "\n".join(content)
        except Exception as e:
            from livebench.utils.logger import log_error
            log_error(f"DOCX extraction failed for {path}", exception=e, context={'path': path})
            raise RuntimeError(f"DOCX extraction failed for {path}: {str(e)}") from e
    
    def _read_xlsx_content(self, path: str) -> str:
        """
        Extract data from XLSX file
        
        Args:
            path: Path to XLSX file
            
        Returns:
            Formatted text representation of spreadsheet data
            
        Raises:
            ImportError: If openpyxl is not installed
            RuntimeError: If XLSX extraction fails
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")
        
        try:
            wb = load_workbook(path, data_only=True)
            
            content = []
            content.append(f"[Excel Workbook - {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}]\n")
            
            for sheet_name in wb.sheetnames[:5]:  # Limit to first 5 sheets
                ws = wb[sheet_name]
                content.append(f"\n=== Sheet: {sheet_name} ({ws.max_row} rows × {ws.max_column} cols) ===")
                
                # Get first 20 rows
                for row_idx, row in enumerate(ws.iter_rows(max_row=20, values_only=True), 1):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        content.append(f"Row {row_idx}: {row_text}")
                
                if ws.max_row > 20:
                    content.append(f"... ({ws.max_row - 20} more rows)")
            
            if len(wb.sheetnames) > 5:
                content.append(f"\n... ({len(wb.sheetnames) - 5} more sheets)")
            
            return "\n".join(content)
        except Exception as e:
            from livebench.utils.logger import log_error
            log_error(f"XLSX extraction failed for {path}", exception=e, context={'path': path})
            raise RuntimeError(f"XLSX extraction failed for {path}: {str(e)}") from e
    
    
    def _read_artifacts_with_images(self, artifact_paths: list[str], max_size_kb: int = 2000) -> Dict[str, Dict[str, Union[str, bytes]]]:
        """
        Read artifact file contents with separate handling for images and text
        
        Args:
            artifact_paths: List of artifact file paths
            max_size_kb: Maximum file size to read in KB (default 2MB)
        
        Returns:
            Dictionary mapping file paths to artifact data:
            {
                'path/to/file.png': {'type': 'image', 'format': 'png', 'data': b'...'},
                'path/to/file.docx': {'type': 'text', 'content': '...text...'}
            }
        
        Raises:
            RuntimeError: If file reading fails critically (e.g., PPTX conversion fails)
            FileNotFoundError: If artifact file doesn't exist
        """
        artifacts = {}
        
        for path in artifact_paths:
            file_size = os.path.getsize(path)
            file_ext = os.path.splitext(path)[1].lower()
            
            # Check file size
            if file_size > max_size_kb * 1024:
                from livebench.utils.logger import log_error
                error_msg = f"File too large: {file_size} bytes (>{max_size_kb}KB) - {path}"
                log_error(error_msg, context={'path': path, 'size': file_size})
                raise RuntimeError(error_msg)
            
            if file_size == 0:
                from livebench.utils.logger import log_error
                error_msg = f"Empty file submitted for evaluation: {path}"
                log_error(error_msg, context={'path': path})
                raise ValueError(error_msg)
            
            # Handle images with base64 encoding
            if file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
                with open(path, 'rb') as f:
                    image_data = f.read()
                
                artifacts[path] = {
                    'type': 'image',
                    'format': file_ext[1:],  # Remove leading dot
                    'data': image_data,
                    'size': file_size
                }
            
            # Handle documents with content extraction
            elif file_ext == '.docx':
                content = self._read_docx_content(path)
                if content.startswith("[DOCX file present but extraction failed"):
                    from livebench.utils.logger import log_error
                    log_error(f"DOCX extraction failed: {path}", context={'path': path})
                    raise RuntimeError(f"DOCX extraction failed for {path}: {content}")
                
                artifacts[path] = {
                    'type': 'text',
                    'content': content
                }
            
            elif file_ext == '.xlsx':
                content = self._read_xlsx_content(path)
                if content.startswith("[Excel file present but extraction failed"):
                    from livebench.utils.logger import log_error
                    log_error(f"XLSX extraction failed: {path}", context={'path': path})
                    raise RuntimeError(f"XLSX extraction failed for {path}: {content}")
                
                artifacts[path] = {
                    'type': 'text',
                    'content': content
                }
            
            elif file_ext == '.pptx':
                # Use unified PPTX reader from file_reading.py
                from livebench.tools.productivity.file_reading import read_pptx_as_images
                from livebench.utils.logger import log_error
                
                pptx_images = read_pptx_as_images(Path(path))
                
                if not pptx_images:
                    error_msg = (
                        f"PPTX conversion failed for {path}. "
                        f"Ensure LibreOffice and pdf2image are installed. "
                        f"Install with: sudo apt-get install libreoffice poppler-utils && pip install pdf2image Pillow"
                    )
                    log_error(error_msg, context={'path': path, 'size': file_size})
                    raise RuntimeError(error_msg)
                
                artifacts[path] = {
                    'type': 'pptx_images',
                    'images': pptx_images,
                    'slide_count': len(pptx_images),
                    'size': file_size
                }
            
            elif file_ext == '.pdf':
                # Convert PDF to images (4 pages per combined image)
                from livebench.tools.productivity.file_reading import read_pdf_as_images
                from livebench.utils.logger import log_error
                
                pdf_images = read_pdf_as_images(Path(path))
                
                if not pdf_images:
                    error_msg = (
                        f"PDF conversion failed for {path}. "
                        f"Ensure poppler-utils and pdf2image are installed. "
                        f"Install with: sudo apt-get install poppler-utils && pip install pdf2image Pillow"
                    )
                    log_error(error_msg, context={'path': path, 'size': file_size})
                    raise RuntimeError(error_msg)
                
                # Return PDF as list of image bytes (similar to PPTX)
                approximate_pages = len(pdf_images) * 4
                artifacts[path] = {
                    'type': 'pdf_images',
                    'images': pdf_images,
                    'image_count': len(pdf_images),
                    'approximate_pages': approximate_pages,
                    'size': file_size
                }
            
            else:
                # Try to read as text
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        artifacts[path] = {
                            'type': 'text',
                            'content': f.read()
                        }
                except UnicodeDecodeError:
                    from livebench.utils.logger import log_error
                    error_msg = f"Unsupported binary file type: {file_ext} - {path}"
                    log_error(error_msg, context={'path': path, 'ext': file_ext})
                    raise RuntimeError(error_msg)
        
        return artifacts
    
    def _build_multimodal_evaluation_content(
        self,
        meta_prompt: Dict,
        task: Dict,
        artifact_data: Dict[str, Dict],
        missing_artifacts: list[str],
        description: str
    ) -> List[Dict[str, Union[str, Dict]]]:
        """
        Build multimodal evaluation content with text and images
        
        Returns:
            List of content blocks for OpenAI's multimodal API
        """
        # Extract key components from meta-prompt
        evaluation_prompt = meta_prompt.get('evaluation_prompt', '')
        rubric = meta_prompt.get('evaluation_rubric', {})
        
        # Build the text portion
        text_content = f"""# TASK EVALUATION REQUEST

## Category: {meta_prompt.get('category', 'Unknown')}

## Evaluation Guidelines:
{evaluation_prompt}

## Task Prompt (Original Assignment):
{task.get('prompt', 'N/A')}

## Task Metadata:
- Task ID: {task.get('task_id', 'N/A')}
- Sector: {task.get('sector', 'N/A')}
- Occupation: {task.get('occupation', 'N/A')}
- Reference Files: {', '.join(task.get('reference_files', [])) or 'None'}

## Agent's Description:
{description or 'No description provided'}

## Submitted Artifacts:

"""
        
        # Add text artifacts
        for path, artifact in artifact_data.items():
            if artifact['type'] == 'text':
                text_content += f"\n### File: {os.path.basename(path)}\n```\n{artifact['content']}\n```\n\n"
            elif artifact['type'] == 'image':
                text_content += f"\n### Image: {os.path.basename(path)} ({artifact['format']}, {artifact['size']} bytes)\n[See image below]\n\n"
            elif artifact['type'] == 'pptx_images':
                text_content += f"\n### PowerPoint: {os.path.basename(path)} ({artifact['slide_count']} slides)\n[See slide images below]\n\n"
            elif artifact['type'] == 'pdf_images':
                text_content += f"\n### PDF: {os.path.basename(path)} (~{artifact['approximate_pages']} pages in {artifact['image_count']} combined images)\n[See PDF pages below - 4 pages per image]\n\n"
        
        if missing_artifacts:
            text_content += f"\n## Missing Artifacts:\n"
            for path in missing_artifacts:
                text_content += f"- {path}\n"
        
        text_content += f"""

---

Please evaluate this work according to the rubric above. Output your evaluation in this format:

**OVERALL SCORE:** [0-10]

**DIMENSION SCORES:**
[List dimension scores from rubric]

**KEY FINDINGS:**
[2-3 bullet points on what worked / didn't work]

**FEEDBACK:**
[1-2 paragraph explanation]

**TOP IMPROVEMENTS NEEDED:**
[Numbered list of 3 specific improvements]
"""
        
        # Build multimodal content array
        content = [{"type": "text", "text": text_content}]
        
        # Add images and PPTX slides
        for path, artifact in artifact_data.items():
            if artifact['type'] == 'image':
                # Convert to base64
                image_base64 = base64.b64encode(artifact['data']).decode('utf-8')
                
                # Determine MIME type
                format_to_mime = {
                    'png': 'image/png',
                    'jpg': 'image/jpeg',
                    'jpeg': 'image/jpeg',
                    'gif': 'image/gif',
                    'webp': 'image/webp'
                }
                mime_type = format_to_mime.get(artifact['format'], 'image/png')
                
                # Add image content block
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime_type};base64,{image_base64}",
                        "detail": "high"  # Request high-resolution analysis
                    }
                })
            
            elif artifact['type'] == 'pptx_images':
                # Add each slide as a separate image
                for i, slide_img_bytes in enumerate(artifact['images'], 1):
                    slide_base64 = base64.b64encode(slide_img_bytes).decode('utf-8')
                    
                    content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{slide_base64}",
                            "detail": "high"
                        }
                    })
            
            elif artifact['type'] == 'pdf_images':
                # Add each combined PDF image (4 pages per image)
                for i, pdf_img_bytes in enumerate(artifact['images'], 1):
                    pdf_base64 = base64.b64encode(pdf_img_bytes).decode('utf-8')
                    
                    content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{pdf_base64}",
                            "detail": "high"
                        }
                    })
        
        return content

    def _build_evaluation_prompt(
        self,
        meta_prompt: Dict,
        task: Dict,
        artifact_contents: Dict[str, str],
        missing_artifacts: list[str],
        description: str
    ) -> str:
        """Build the evaluation prompt for the LLM"""
        
        # Extract key components from meta-prompt
        evaluation_prompt = meta_prompt.get('evaluation_prompt', '')
        rubric = meta_prompt.get('evaluation_rubric', {})
        
        # Build the complete evaluation request
        prompt = f"""# TASK EVALUATION REQUEST

## Category: {meta_prompt.get('category', 'Unknown')}

## Evaluation Guidelines:
{evaluation_prompt}

## Task Prompt (Original Assignment):
{task.get('prompt', 'N/A')}

## Task Metadata:
- Task ID: {task.get('task_id', 'N/A')}
- Sector: {task.get('sector', 'N/A')}
- Occupation: {task.get('occupation', 'N/A')}
- Reference Files: {', '.join(task.get('reference_files', [])) or 'None'}

## Agent's Description:
{description or 'No description provided'}

## Submitted Artifacts:

"""
        
        # Add artifact contents
        if artifact_contents:
            for path, content in artifact_contents.items():
                prompt += f"\n### File: {os.path.basename(path)}\n```\n{content}\n```\n\n"
        else:
            prompt += "⚠️ No artifacts were submitted or found.\n\n"
        
        # Add missing artifacts warning
        if missing_artifacts:
            prompt += f"\n⚠️ **MISSING ARTIFACTS:** {', '.join(missing_artifacts)}\n\n"
        
        # Add evaluation instructions
        prompt += f"""
---

## YOUR EVALUATION TASK:

1. Review the task requirements carefully
2. Assess completeness, correctness, quality, and domain standards per the rubric
3. Apply the CRITICAL POLICY: If ANY required artifacts are missing or work is severely incomplete, score must be 0-2
4. Provide a structured evaluation

## OUTPUT FORMAT (REQUIRED):

**OVERALL SCORE:** [0-10 integer]

**DIMENSION SCORES:**
- Completeness: [0-10]
- Correctness: [0-10]
- Quality: [0-10]
- Domain Standards: [0-10]

**KEY FINDINGS:**
- [Finding 1]
- [Finding 2]
- [Finding 3]

**FEEDBACK:**
[2-3 sentences of constructive feedback]

**TOP IMPROVEMENTS NEEDED:**
1. [Improvement 1]
2. [Improvement 2]
3. [Improvement 3]

Remember: The scoring scale is 0-10, where:
- 0-2: Unacceptable (missing files/incomplete)
- 3-4: Poor (major issues)
- 5-6: Acceptable (notable gaps)
- 7-8: Good (minor issues)
- 9-10: Excellent (complete, accurate, professional)
"""
        
        return prompt

    def _extract_score(self, evaluation_text: str) -> float:
        """
        Extract numerical score from LLM evaluation response

        Args:
            evaluation_text: Full evaluation text from LLM

        Returns:
            Score as float (0-10 scale)
        """
        import re
        
        # Look for "OVERALL SCORE: X" pattern
        patterns = [
            r'OVERALL SCORE:\s*(\d+(?:\.\d+)?)',
            r'Overall Score:\s*(\d+(?:\.\d+)?)',
            r'Score:\s*(\d+(?:\.\d+)?)/10',
            r'Final Score:\s*(\d+(?:\.\d+)?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, evaluation_text, re.IGNORECASE)
            if match:
                score = float(match.group(1))
                # Clamp to 0-10 range
                return max(0.0, min(10.0, score))
        
        # If no score found, look for any number in first 200 chars
        first_part = evaluation_text[:200]
        numbers = re.findall(r'\b(\d+(?:\.\d+)?)\b', first_part)
        
        if numbers:
            score = float(numbers[0])
            if 0 <= score <= 10:
                return score
        
        # Default to 5.0 if no score found
        print("⚠️ Could not extract score from evaluation, defaulting to 5.0")
        return 5.0

    # REMOVED: Fallback evaluation method
    # System now requires LLM evaluation to ensure quality and consistency
    # Errors will propagate if LLM evaluation fails


if __name__ == "__main__":
    """Test the LLM evaluator"""
    
    # Example task
    test_task = {
        "task_id": "test-001",
        "occupation": "Software_Developers",
        "sector": "Technology",
        "prompt": "Create a simple Python function to calculate Fibonacci numbers.",
        "reference_files": []
    }
    
    # Example artifact (would be actual file in practice)
    test_artifact = "/tmp/test_artifact.py"
    with open(test_artifact, "w") as f:
        f.write("""def fibonacci(n):
    if n <= 1:
        return n
    return fibonacci(n-1) + fibonacci(n-2)

# Test
print(fibonacci(10))
""")
    
    # Evaluate
    evaluator = LLMEvaluator()
    score, feedback, payment = evaluator.evaluate_artifact(
        task=test_task,
        artifact_paths=[test_artifact],
        description="Implemented recursive Fibonacci function"
    )
    
    print(f"\n{'='*60}")
    print(f"EVALUATION RESULTS")
    print(f"{'='*60}")
    print(f"Score: {score:.2f} (0.0-1.0 scale)")
    print(f"Payment: ${payment:.2f}")
    print(f"\nFeedback:\n{feedback}")
    print(f"{'='*60}\n")
    
    # Cleanup
    os.remove(test_artifact)

